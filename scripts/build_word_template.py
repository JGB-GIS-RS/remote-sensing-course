from pathlib import Path
import io, zipfile
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw
from lxml import etree

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'template'; OUT.mkdir(exist_ok=True)
DOCX=OUT/'Plantilla_Oficial_Serie_Tecnica_Teledeteccion_v1.0.docx'
DOTX=OUT/'Plantilla_Oficial_Serie_Tecnica_Teledeteccion_v1.0.dotx'
LOGO=ROOT/'.template_logo.png'
GREEN='00843D'; DARK='303238'; GRAY='666B73'; LIGHT='F2F4F3'; RULE='D3D7D5'; FONT='Carlito'

def font(r,s=None,b=None,i=None,c=None):
 r.font.name=FONT; r._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
 if s:r.font.size=Pt(s)
 if b is not None:r.bold=b
 if i is not None:r.italic=i
 if c:r.font.color.rgb=RGBColor.from_string(c)

def border_p(p,color=RULE,sz='6'):
 pp=p._p.get_or_add_pPr(); pb=pp.find(qn('w:pBdr'))
 if pb is None:pb=OxmlElement('w:pBdr');pp.append(pb)
 e=OxmlElement('w:bottom'); e.set(qn('w:val'),'single');e.set(qn('w:sz'),sz);e.set(qn('w:space'),'1');e.set(qn('w:color'),color);pb.append(e)

def cmarg(c,t=50,s=70,b=50,e=70):
 pr=c._tc.get_or_add_tcPr(); m=pr.first_child_found_in('w:tcMar')
 if m is None:m=OxmlElement('w:tcMar');pr.append(m)
 for k,v in [('top',t),('start',s),('bottom',b),('end',e)]:
  x=OxmlElement('w:'+k);x.set(qn('w:w'),str(v));x.set(qn('w:type'),'dxa');m.append(x)

def cborder(c,**edges):
 pr=c._tc.get_or_add_tcPr(); bs=pr.first_child_found_in('w:tcBorders')
 if bs is None:bs=OxmlElement('w:tcBorders');pr.append(bs)
 for edge,spec in edges.items():
  x=OxmlElement('w:'+edge)
  for k,v in spec.items():x.set(qn('w:'+k),str(v))
  bs.append(x)

def shade(c,fill):
 pr=c._tc.get_or_add_tcPr();x=OxmlElement('w:shd');x.set(qn('w:fill'),fill);pr.append(x)

def page_field(p):
 r=p.add_run();a=OxmlElement('w:fldChar');a.set(qn('w:fldCharType'),'begin');b=OxmlElement('w:instrText');b.set(qn('xml:space'),'preserve');b.text=' PAGE ';c=OxmlElement('w:fldChar');c.set(qn('w:fldCharType'),'end');r._r.extend([a,b,c])

def logo():
 im=Image.new('RGBA',(180,180),(0,0,0,0));d=ImageDraw.Draw(im)
 d.rounded_rectangle((20,20,160,160),radius=12,fill=(0,93,45,255));d.polygon([(20,20),(90,20),(20,90)],fill=(0,112,56,255));d.polygon([(160,20),(160,90),(90,20)],fill=(0,74,35,255));d.rounded_rectangle((72,55,108,115),radius=17,fill='white');d.rectangle((86,113,94,132),fill='white');d.rounded_rectangle((76,132,104,140),radius=4,fill='white');im.resize((90,90),Image.Resampling.LANCZOS).save(LOGO,optimize=True)

def style(doc,name,size=9.5,color=DARK,bold=False,italic=False,before=0,after=0,line=1.0):
 st=doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH);st.base_style=doc.styles['Normal'];st.font.name=FONT;st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT);st.font.size=Pt(size);st.font.color.rgb=RGBColor.from_string(color);st.font.bold=bold;st.font.italic=italic;pf=st.paragraph_format;pf.space_before=Pt(before);pf.space_after=Pt(after);pf.line_spacing=line

def guide(doc,text):return doc.add_paragraph(text,style='ITT Guide')
def head(doc,text,n=1):return doc.add_paragraph(text,style='ITT Heading 1' if n==1 else 'ITT Heading 2')

def fig(doc,caption):
 t=doc.add_table(1,1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;c=t.cell(0,0);c.width=Cm(16.8);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;shade(c,LIGHT);cmarg(c,430,70,430,70);cborder(c,top={'val':'single','sz':'6','color':'B7BCB9'},bottom={'val':'single','sz':'6','color':'B7BCB9'},start={'val':'single','sz':'6','color':'B7BCB9'},end={'val':'single','sz':'6','color':'B7BCB9'});p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(p.add_run('[INSERTE AQUÍ LA FIGURA SI ES PERTINENTE]'),9,True,None,GRAY);doc.add_paragraph('Figura X. '+caption,style='ITT Caption')

def table(doc):
 doc.add_paragraph('Tabla X. [Título preciso y autosuficiente de la tabla].',style='ITT Table Title');t=doc.add_table(4,2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
 rows=[('Elemento','Especificación / valor'),('[Dato o producto]','[Descripción]'),('[Periodo / escala]','[Valor]'),('[Fuente]','[Fuente oficial / DOI / URL de recuperación]')]
 for i,row in enumerate(t.rows):
  for j,c in enumerate(row.cells):
   c.width=Cm(8.5 if j==0 else 8.3);cmarg(c,30,80,30,80)
   if i==0:shade(c,'EAF2ED');cborder(c,top={'val':'single','sz':'12','color':GREEN},bottom={'val':'single','sz':'6','color':GREEN})
   else:cborder(c,bottom={'val':'single','sz':'4','color':'D8DCDA'})
   font(c.paragraphs[0].add_run(rows[i][j]),8.7,i==0,None,DARK)
 doc.add_paragraph('Nota: [Incluya abreviaturas, fuente o aclaraciones únicamente cuando sean necesarias].',style='ITT Note')

def build():
 logo();d=Document();s=d.sections[0];s.top_margin=Cm(1.35);s.bottom_margin=Cm(1.35);s.left_margin=Cm(1.85);s.right_margin=Cm(1.55);s.header_distance=Cm(.55);s.footer_distance=Cm(.55);s.different_first_page_header_footer=True
 n=d.styles['Normal'];n.font.name=FONT;n._element.rPr.rFonts.set(qn('w:eastAsia'),FONT);n.font.size=Pt(9.5);n.font.color.rgb=RGBColor.from_string(DARK);n.paragraph_format.space_after=Pt(3);n.paragraph_format.line_spacing=1
 style(d,'ITT Title',17.5,DARK,True,False,2,9);style(d,'ITT Authors',10.5,DARK,False,False,0,3);style(d,'ITT Meta',8.8,GRAY,False,False,0,1);style(d,'ITT Guide',9.2,GRAY,False,True,0,6,1.05);style(d,'ITT Heading 1',12.5,GREEN,True,False,8,4);style(d,'ITT Heading 2',10.3,DARK,True,False,6,3);style(d,'ITT Caption',8.1,DARK,False,False,3,5);style(d,'ITT Table Title',8.1,DARK,False,False,4,2);style(d,'ITT Note',8,GRAY,False,False,2,5);style(d,'ITT References',7.7,DARK,False,False,0,1)
 ln=OxmlElement('w:lnNumType');ln.set(qn('w:countBy'),'1');ln.set(qn('w:restart'),'continuous');ln.set(qn('w:distance'),'260');s._sectPr.insert(0,ln)
 h=s.first_page_header;t=h.add_table(1,2,Cm(8.2));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;t.columns[0].width=Cm(6.5);t.columns[1].width=Cm(1.2);l,r=t.rows[0].cells;l.width=Cm(6.5);r.width=Cm(1.2);cmarg(l,0,0,0,120);cmarg(r,0,0,0,0);p=l.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;font(p.add_run('SERIE TÉCNICA DE TELEDETECCIÓN'),9.2,True,None,GREEN);p=l.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;font(p.add_run('Informe Técnico No. [XX]'),8.8,None,None,GRAY);p=r.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(LOGO),width=Cm(1));p=h.add_paragraph();border_p(p,GREEN,'10')
 p=s.header.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(p.add_run('Informe Técnico No. [XX]    ·    Serie Técnica de Teledetección'),8.2,None,None,GRAY);border_p(p,RULE,'5')
 for f in (s.footer,s.first_page_footer):p=f.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;page_field(p)
 p=d.add_paragraph('Universidad del Quindío · Facultad de Ingeniería · Programa de Ingeniería Topográfica y Geomática',style='ITT Meta');border_p(p,GREEN,'8');d.add_paragraph('[Título científico e informativo del informe técnico]',style='ITT Title');d.add_paragraph('[Nombre Apellido] · [Nombre Apellido]',style='ITT Authors');d.add_paragraph('Correos: [correo1@uq.edu.co] · [correo2@uq.edu.co]',style='ITT Meta');d.add_paragraph('Teledetección · Semestre [AAAA-X] · [día de mes de año]',style='ITT Meta');p=d.add_paragraph();border_p(p,GREEN,'8')
 guide(d,'TEXTO GUÍA EN GRIS: reemplácelo o elimínelo antes de entregar. Los grandes encabezados de la Serie deben conservarse; las subsecciones se adaptan a la TAREA.')
 head(d,'Resumen',2);guide(d,'[150–200 palabras aprox. Sintetice objetivo/problema, datos y método esencial, resultados principales y conclusión o decisión técnica. No incluya información que no aparezca en el cuerpo del informe.]');p=d.add_paragraph();font(p.add_run('Palabras clave: '),9.5,True);font(p.add_run('[4–6 términos o expresiones breves, separados por punto y coma].'),9.5)
 head(d,'1. Introducción');guide(d,'[Plantee el problema técnico, el conocimiento previo estrictamente necesario y la motivación metodológica. Finalice con el objetivo específico del informe. Evite un marco teórico extenso.]')
 head(d,'2. Área de estudio y datos');head(d,'2.1. Área de estudio',2);guide(d,'[Describa solo las características territoriales necesarias para interpretar el problema y los resultados: localización, escala, relieve, clima/estacionalidad, coberturas u otros rasgos pertinentes.]');fig(d,'[Localización o evidencia espacial pertinente. Explique paneles, fecha, unidades, simbología y fuente cuando corresponda].');head(d,'2.2. Datos',2);guide(d,'[Identifique misión/plataforma, sensor, producto/colección, nivel de procesamiento, fecha o periodo, resoluciones relevantes, magnitud física, factores de escala, QA/correcciones y fuente de acceso cuando corresponda. Justifique la pertinencia de los datos.]');table(d)
 head(d,'3. Metodología');guide(d,'[Describa operaciones y decisiones, no secuencias de botones. Documente criterios de selección, preprocesamiento realmente aplicado, parámetros, variables/transformaciones, procedimiento de análisis, validación y métricas cuando correspondan. Cree subsecciones 3.1, 3.2, ... según la TAREA.]');d.add_page_break();head(d,'3.1. [Subsección metodológica pertinente]',2);guide(d,'[Método, parámetros, criterios y decisiones reproducibles. Si utiliza ecuaciones, insértelas con el editor de ecuaciones y numérelas cuando deban referenciarse.]')
 head(d,'4. Resultados');guide(d,'[Presente qué se obtuvo: valores, patrones, tendencias y comparaciones. No repita fila por fila una tabla ni describa mecánicamente una figura. No use “significativo” sin soporte inferencial.]');head(d,'4.1. [Resultado principal / evidencia 1]',2);guide(d,'[Reporte la evidencia principal de manera compacta. Añada únicamente las subsecciones necesarias para organizar resultados distintos.]');fig(d,'[Resultado o comparación principal. El pie debe permitir interpretar la figura de manera autónoma].')
 head(d,'5. Discusión');guide(d,'[Interprete qué significan los resultados, contraste con literatura pertinente, explique mecanismos posibles con cautela y analice compromisos metodológicos. No repita los Resultados.]');head(d,'5.1. Alcance de los resultados e incertidumbre metodológica',2);guide(d,'[Declare supuestos, fuentes de incertidumbre, limitaciones del dato/método y condiciones de validez. Explique cómo cada limitación afecta la interpretación o transferibilidad.]')
 head(d,'6. Recomendación técnica');guide(d,'[Convierta la evidencia en una decisión profesional: qué hacer, con qué datos/método/condiciones y bajo qué límites. Evite recomendaciones universales que el estudio no permita sostener.]');head(d,'7. Conclusiones');guide(d,'[Responda al objetivo con hallazgos declarativos y sintéticos. No introduzca resultados ni bibliografía nuevos y no convierta esta sección en otra recomendación.]');head(d,'Disponibilidad de datos y código',2);guide(d,'[Indique procedencia de datos originales y suministrados, productos derivados, scripts/notebooks, software/versiones relevantes, ubicación del material reproducible y restricciones de acceso. Declare el uso de IA generativa cuando corresponda según la Guía para autores.]')
 head(d,'Referencias');d.add_paragraph('[1] A. Autor, B. Autor, “Título del artículo,” Abrev. Revista, vol. X, no. X, pp. xx–xx, año, doi: xx.xxxx/xxxxx.',style='ITT References');d.add_paragraph('[2] Organización, Título del manual o documentación oficial, versión, año. [Datos de recuperación cuando correspondan].',style='ITT References');guide(d,'Antes de entregar: elimine todo el texto guía en gris, verifique el límite de 8 páginas, conserve la estructura 1–7 y revise la Guía para autores v1.0.')
 cp=d.core_properties;cp.title='Plantilla Oficial · Serie Técnica de Teledetección v1.0';cp.subject='Plantilla para Informes Técnicos de Teledetección';cp.author='Universidad del Quindío · Programa de Ingeniería Topográfica y Geomática';cp.last_modified_by=cp.author;cp.comments='Plantilla oficial v1.0';d.save(DOCX)
 CT='{http://schemas.openxmlformats.org/package/2006/content-types}'
 with zipfile.ZipFile(DOCX) as zi,zipfile.ZipFile(DOTX,'w',zipfile.ZIP_DEFLATED,compresslevel=9) as zo:
  for item in zi.infolist():
   data=zi.read(item.filename)
   if item.filename=='[Content_Types].xml':
    root=etree.fromstring(data)
    for ov in root.findall(CT+'Override'):
     if ov.get('PartName')=='/word/document.xml':ov.set('ContentType','application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml')
    data=etree.tostring(root,xml_declaration=True,encoding='UTF-8',standalone='yes')
   zo.writestr(item,data)
 LOGO.unlink(missing_ok=True)
if __name__=='__main__':build()
